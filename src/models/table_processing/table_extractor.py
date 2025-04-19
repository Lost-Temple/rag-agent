import pandas as pd
from docx.api import Document
import camelot
from langchain.schema import Document as LangchainDocument

class TableExtractor:
    """表格数据提取器，专门处理文档中的表格"""
    
    def extract_tables_from_pdf(self, pdf_path):
        """从PDF中提取表格并转换为结构化文本"""
        # 使用camelot提取表格（更精确）
        tables = camelot.read_pdf(pdf_path, pages='all', flavor='lattice')
        
        extracted_docs = []
        for i, table in enumerate(tables):
            # 将表格转换为DataFrame
            df = table.df
            
            # 生成表格的文本描述
            table_text = f"表格{i+1}内容:\n"
            # 添加列名
            table_text += "列标题: " + " | ".join(df.iloc[0].tolist()) + "\n"
            
            # 添加每行数据，并进行自然语言描述
            for idx, row in df.iloc[1:].iterrows():
                row_values = row.tolist()
                table_text += f"行{idx}: " + " | ".join(row_values) + "\n"
                
                # 为每行生成自然语言描述
                col_names = df.iloc[0].tolist()
                row_desc = "该行数据表示: "
                for j, val in enumerate(row_values):
                    if j < len(col_names):
                        row_desc += f"{col_names[j]}为{val}, "
                table_text += row_desc.rstrip(", ") + "\n"
            
            # 创建文档对象
            metadata = {"source": pdf_path, "content_type": "table", "table_index": i}
            doc = LangchainDocument(page_content=table_text, metadata=metadata)
            extracted_docs.append(doc)
            
        return extracted_docs
    
    def extract_tables_from_docx(self, docx_path):
        """从Word文档中提取表格并转换为结构化文本"""
        doc = Document(docx_path)
        
        extracted_docs = []
        for i, table in enumerate(doc.tables):
            # 提取表格数据到DataFrame
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
                
            if not data:
                continue
                
            df = pd.DataFrame(data[1:], columns=data[0] if data else None)
            
            # 生成表格的文本描述（与PDF处理类似）
            table_text = f"表格{i+1}内容:\n"
            # 添加列名
            if len(data) > 0:
                table_text += "列标题: " + " | ".join(data[0]) + "\n"
            
            # 添加每行数据，并进行自然语言描述
            for idx, row in enumerate(data[1:]):
                table_text += f"行{idx+1}: " + " | ".join(row) + "\n"
                
                # 为每行生成自然语言描述
                if len(data) > 0:
                    col_names = data[0]
                    row_desc = "该行数据表示: "
                    for j, val in enumerate(row):
                        if j < len(col_names):
                            row_desc += f"{col_names[j]}为{val}, "
                    table_text += row_desc.rstrip(", ") + "\n"
            
            # 创建文档对象
            metadata = {"source": docx_path, "content_type": "table", "table_index": i}
            doc = LangchainDocument(page_content=table_text, metadata=metadata)
            extracted_docs.append(doc)
            
        return extracted_docs